import io
from rest_framework import generics, status, views
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib.auth import get_user_model
from django.core.mail import send_mail
from django.conf import settings
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .serializers import UserSerializer, RegisterSerializer
from .models import PasswordResetToken

User = get_user_model()


class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    permission_classes = [AllowAny]
    serializer_class = RegisterSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        user = serializer.save()
        refresh = RefreshToken.for_user(user)
        return Response({
            'user': UserSerializer(user).data,
            'tokens': {
                'refresh': str(refresh),
                'access': str(refresh.access_token),
            }
        }, status=status.HTTP_201_CREATED)


class UserProfileView(generics.RetrieveUpdateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = UserSerializer

    def get_object(self):
        return self.request.user


class SyncProgressView(views.APIView):
    """Sync user progress to/from database"""
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """Save progress to DB"""
        user = request.user
        # Validate ranges to prevent manipulation
        points = request.data.get('points', user.points)
        tasks = request.data.get('tasks_completed', user.tasks_completed)
        hours = request.data.get('study_hours', user.study_hours)
        # Only allow increases up to reasonable limits
        user.points = max(0, min(int(points), 100000))
        user.level = max(1, min(int(request.data.get('level', user.level)), 100))
        user.streak = max(0, min(int(request.data.get('streak', user.streak)), 3650))
        user.tasks_completed = max(0, min(int(tasks), 10000))
        user.study_hours = max(0.0, min(float(hours), 10000.0))
        user.badges = request.data.get('badges', user.badges)
        user.roadmap_tasks = request.data.get('roadmap_tasks', user.roadmap_tasks)
        user.save(update_fields=['points', 'level', 'streak', 'tasks_completed', 'study_hours', 'badges', 'roadmap_tasks'])
        return Response({'message': 'Progress saved'})

    def get(self, request):
        """Load progress from DB"""
        user = request.user
        return Response({
            'points': user.points,
            'level': user.level,
            'streak': user.streak,
            'tasks_completed': user.tasks_completed,
            'study_hours': user.study_hours,
            'badges': user.badges,
            'roadmap_tasks': user.roadmap_tasks,
        })


class GenerateResumeView(views.APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        career_title    = request.data.get('career_title', '')
        career_category = request.data.get('career_category', '')
        career_skills   = request.data.get('career_skills', [])
        career_education= request.data.get('career_education', '')
        career_salary   = request.data.get('career_salary', '')
        phone           = request.data.get('phone', '')
        location        = request.data.get('location', '')
        points          = request.data.get('points', 0)
        level           = request.data.get('level', 1)
        tasks_completed = request.data.get('tasks_completed', 0)
        study_hours     = request.data.get('study_hours', 0)
        user            = request.user
        skills_str      = ', '.join(career_skills) if career_skills else 'Not specified'

        # --- Generate AI content via Gemini ---
        prompt = f"""Write a professional resume for the following person. Return ONLY the resume content in this EXACT format with these section headers:

PROFESSIONAL SUMMARY
[3-4 sentences tailored to {career_title}]

CAREER OBJECTIVE
[2-3 sentences about career goals]

KEY SKILLS
- [skill 1]
- [skill 2]
- [skill 3]
- [skill 4]
- [skill 5]

EDUCATION
[Degree/qualification based on {career_education}]

PROJECTS & ACHIEVEMENTS
- Completed {tasks_completed} structured learning tasks on AI Career Compass
- Accumulated {round(study_hours, 1)} hours of focused study
- Earned {points} XP — Level {level} learner
- [1-2 relevant project ideas for {career_title}]

CERTIFICATIONS & TRAINING
- AI Career Compass — {career_title} Learning Path ({career_category})
- [1-2 relevant certifications for {career_title}]

ADDITIONAL INFORMATION
- Target Salary: {career_salary}
- Available for: Full-time, {career_category} roles

Person: {user.username}, {user.email}, {phone or 'N/A'}, {location or 'N/A'}
Career: {career_title} ({career_category})
Skills: {skills_str}"""

        try:
            from careers.llm_service import LLMService
            llm = LLMService()
            ai_content = llm._call_gemini(prompt, temperature=0.6)
        except Exception as e:
            ai_content = None

        # --- Build DOCX ---
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        # Helper: add heading
        def add_section_heading(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)  # blue
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '1a56db')
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_body(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)

        def add_bullet(text):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            run.font.size = Pt(10)

        # ── HEADER ──
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(user.username.upper())
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

        contact_parts = [user.email]
        if phone:    contact_parts.append(phone)
        if location: contact_parts.append(location)
        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)

        if career_title:
            role_para = doc.add_paragraph(career_title)
            role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in role_para.runs:
                run.bold = True
                run.font.size = Pt(11)

        doc.add_paragraph()  # spacer

        # ── AI SECTIONS ──
        if ai_content:
            for line in ai_content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                upper = line.upper()
                if any(upper.startswith(h) for h in [
                    'PROFESSIONAL SUMMARY', 'CAREER OBJECTIVE', 'KEY SKILLS',
                    'EDUCATION', 'PROJECTS', 'CERTIFICATIONS', 'ADDITIONAL'
                ]):
                    current_section = upper
                    add_section_heading(line)
                elif line.startswith('- ') or line.startswith('• '):
                    add_bullet(line.lstrip('-• ').strip())
                else:
                    add_body(line)
        else:
            # Fallback if AI fails
            add_section_heading('Professional Summary')
            add_body(f'Motivated {career_title} candidate with strong foundational knowledge and a commitment to continuous learning.')
            add_section_heading('Key Skills')
            for skill in career_skills:
                add_bullet(skill)
            add_section_heading('Education')
            add_body(career_education or 'Relevant degree/qualification')
            add_section_heading('Achievements')
            add_bullet(f'{tasks_completed} learning tasks completed')
            add_bullet(f'{round(study_hours, 1)} study hours logged')
            add_bullet(f'{points} XP earned — Level {level}')

        # ── Save to buffer ──
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{user.username.replace(' ', '_')}_Resume.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response


class AdminUsersView(views.APIView):
    """Panel demo endpoint — shows all registered users (no auth required)"""
    permission_classes = [AllowAny]

    def get(self, request):
        users = User.objects.all().order_by('date_joined').values(
            'id', 'username', 'email', 'points', 'level',
            'streak', 'tasks_completed', 'study_hours', 'date_joined'
        )
        return Response({
            'total_users': User.objects.count(),
            'users': list(users)
        })
    permission_classes = [AllowAny]

    def post(self, request):
        email = request.data.get('email', '').strip().lower()
        if not email:
            return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return Response({'error': 'not_registered'}, status=status.HTTP_404_NOT_FOUND)

        # Invalidate old tokens
        PasswordResetToken.objects.filter(user=user, used=False).update(used=True)

        # Create new token
        token_obj = PasswordResetToken.objects.create(user=user)
        reset_url = f"{settings.FRONTEND_URL}/reset-password?token={token_obj.token}"

        try:
            send_mail(
                subject='AI Career Compass — Password Reset',
                message=f"""Hi {user.username},

You requested a password reset for your AI Career Compass account.

Click the link below to reset your password (valid for 1 hour):
{reset_url}

If you did not request this, please ignore this email.

— AI Career Compass Team""",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[user.email],
                fail_silently=False,
            )
        except Exception as e:
            print(f"Email send error: {e}")
            return Response({'error': 'Failed to send email. Please try again.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response({'message': 'If this email is registered, a reset link has been sent.'})


class ResetPasswordView(views.APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        token_str = request.data.get('token', '').strip()
        new_password = request.data.get('password', '').strip()

        if not token_str or not new_password:
            return Response({'error': 'Token and password are required'}, status=status.HTTP_400_BAD_REQUEST)

        if len(new_password) < 6:
            return Response({'error': 'Password must be at least 6 characters'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            token_obj = PasswordResetToken.objects.get(token=token_str)
        except (PasswordResetToken.DoesNotExist, ValueError):
            return Response({'error': 'Invalid or expired reset link'}, status=status.HTTP_400_BAD_REQUEST)

        if not token_obj.is_valid():
            return Response({'error': 'Reset link has expired. Please request a new one.'}, status=status.HTTP_400_BAD_REQUEST)

        user = token_obj.user
        user.set_password(new_password)
        user.save()

        token_obj.used = True
        token_obj.save()

        return Response({'message': 'Password reset successfully. You can now log in.'})
